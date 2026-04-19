import io
import streamlit as st
import pandas as pd
import sys
from pathlib import Path
from datetime import date, datetime

sys.path.insert(0, str(Path(__file__).parent.parent))
from auth.auth_ui import require_login
from db.database import query, execute, get_connection
from utils.helpers import sidebar_quick_stats

require_login(tool="tracker")
sidebar_quick_stats()

css_path = Path(__file__).parent.parent / "assets" / "style.css"
if css_path.exists():
    st.markdown(f"<style>{css_path.read_text()}</style>", unsafe_allow_html=True)

st.markdown("""
<div style="margin-bottom:20px;">
    <h1 style="margin:0; font-size:1.8rem;">Settings</h1>
    <p style="color:#64748b; margin:4px 0 0 0; font-size:0.95rem;">Manage roles, rates, and project data</p>
</div>
""", unsafe_allow_html=True)


# ─── Report generation helpers ────────────────────────────────────────────────

def _fmt_price(v):
    """Format a price value: no decimals when whole number, else 2dp."""
    if v == int(v):
        return f"${int(v):,}"
    return f"${v:,.2f}"


def _set_cell_text(cell, text):
    """Replace all runs in a table cell's first paragraph with a single text run."""
    from docx.oxml.ns import qn
    para = cell.paragraphs[0]
    p = para._p
    for r_elem in p.findall(qn("w:r")):
        p.remove(r_elem)
    if text:
        para.add_run(str(text))


def _generate_docx(header_data, labor_row, img_rows):
    """Fill the Options Invoice template and return DOCX bytes."""
    from docx import Document
    template_path = Path(__file__).parent.parent / "assets" / "Options_Invoice_Template.docx"
    doc = Document(str(template_path))

    # Table 0 — header fields
    t0 = doc.tables[0]
    _set_cell_text(t0.rows[1].cells[1], header_data["contract_event"])
    _set_cell_text(t0.rows[2].cells[1], header_data["submission_date"])

    # Table 1 — options breakdown
    t1 = doc.tables[1]
    all_rows = [labor_row] + img_rows + [{"name": "Nothing Follows", "qty": "", "price": "", "desc": ""}]

    for i, rd in enumerate(all_rows):
        row_idx = i + 1  # row 0 is the header
        if row_idx >= len(t1.rows):
            # Copy last data row XML to add a new row
            from copy import deepcopy
            last_tr = t1.rows[-1]._tr
            new_tr = deepcopy(last_tr)
            t1._tbl.append(new_tr)

        tr = t1.rows[row_idx]
        _set_cell_text(tr.cells[0], rd["name"])
        _set_cell_text(tr.cells[1], str(rd["qty"]) if rd["qty"] != "" else "")
        _set_cell_text(tr.cells[2], _fmt_price(rd["price"]) if rd["price"] != "" else "")
        _set_cell_text(tr.cells[3], rd.get("desc", ""))

    # Clear leftover empty rows
    clear_from = len(all_rows) + 1
    for i in range(clear_from, len(t1.rows)):
        tr = t1.rows[i]
        for cell in tr.cells:
            _set_cell_text(cell, "")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _generate_pdf(header_data, labor_row, img_rows):
    """Generate a clean PDF version of the invoice using reportlab."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch

    buf = io.BytesIO()
    # Use letter size with generous margins so content fits
    page_w, page_h = letter
    margin = 0.6 * inch
    usable_w = page_w - 2 * margin

    doc = SimpleDocTemplate(
        buf, pagesize=letter,
        topMargin=margin, bottomMargin=margin,
        leftMargin=margin, rightMargin=margin,
    )
    styles = getSampleStyleSheet()
    navy = colors.HexColor("#1f4e79")
    light_grey = colors.HexColor("#f2f2f2")
    border_grey = colors.HexColor("#bbbbbb")

    cell_style = ParagraphStyle("Cell", parent=styles["Normal"], fontSize=8, leading=11, wordWrap="LTR")
    cell_bold = ParagraphStyle("CellBold", parent=cell_style, fontName="Helvetica-Bold")
    cell_white = ParagraphStyle("CellWhite", parent=cell_style, fontName="Helvetica-Bold", textColor=colors.white)

    title_style = ParagraphStyle("T", parent=styles["Title"], fontSize=15, spaceAfter=2, textColor=navy)
    sub_style = ParagraphStyle("S", parent=styles["Normal"], fontSize=8, spaceAfter=10, textColor=colors.grey)
    h2_style = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=11, spaceBefore=12, spaceAfter=6, textColor=navy)

    story = [
        Paragraph("Options Invoice Sheet", title_style),
        Paragraph("(Please complete NLT 24 hrs after AAR submission and please convert to PDF)", sub_style),
    ]

    # Header info table — two equal columns
    col_w = usable_w / 2
    t0_data = [
        [Paragraph("Vendor Name:", cell_bold), Paragraph("Stone Harp Analytics", cell_style)],
        [Paragraph("Contract Number \u2013 Event Name:", cell_bold), Paragraph(header_data["contract_event"], cell_style)],
        [Paragraph("Option Submission Date:", cell_bold), Paragraph(header_data["submission_date"], cell_style)],
        [Paragraph("Does the following options reflect the original options?", cell_bold), Paragraph("Y", cell_style)],
        [Paragraph("If not, please explain:", cell_bold), Paragraph("No Modification Required", cell_style)],
        [Paragraph("Space Component:", cell_bold), Paragraph("To be Filled by Customer", cell_style)],
    ]
    t0 = Table(t0_data, colWidths=[col_w, col_w])
    t0.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, border_grey),
        ("BACKGROUND", (0, 0), (0, -1), light_grey),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(t0)
    story.append(Paragraph("Options Breakdown Assessment", h2_style))

    # Options table — fixed column widths that sum to usable_w
    # Name | Qty | Cost | Description
    c0 = 1.7 * inch
    c1 = 0.65 * inch
    c2 = 1.1 * inch
    c3 = usable_w - c0 - c1 - c2

    hdr_cells = [
        Paragraph("Name", cell_white),
        Paragraph("Quantity", cell_white),
        Paragraph("Cost per unit ($)", cell_white),
        Paragraph("Description / Breakdown", cell_white),
    ]
    t1_data = [hdr_cells]
    t1_data.append([
        Paragraph(labor_row["name"], cell_bold),
        Paragraph(str(labor_row["qty"]), cell_style),
        Paragraph(_fmt_price(labor_row["price"]), cell_style),
        Paragraph(labor_row.get("desc", ""), cell_style),
    ])
    for r in img_rows:
        t1_data.append([
            Paragraph(r["name"], cell_style),
            Paragraph(str(r["qty"]), cell_style),
            Paragraph(_fmt_price(r["price"]), cell_style),
            Paragraph(r.get("desc", ""), cell_style),
        ])
    t1_data.append([
        Paragraph("Nothing Follows", ParagraphStyle("NF", parent=cell_style, textColor=colors.grey, fontName="Helvetica-Oblique")),
        Paragraph("", cell_style), Paragraph("", cell_style), Paragraph("", cell_style),
    ])

    row_count = len(t1_data)
    t1 = Table(t1_data, colWidths=[c0, c1, c2, c3], repeatRows=1)
    t1.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, border_grey),
        ("BACKGROUND", (0, 0), (-1, 0), navy),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("ROWBACKGROUNDS", (0, 1), (-1, row_count - 2), [colors.white, colors.HexColor("#f0f4f8")]),
    ]))
    story.append(t1)

    doc.build(story)
    buf.seek(0)
    return buf.getvalue()


# ─── Page tabs ────────────────────────────────────────────────────────────────

tab_roles, tab_db, tab_report = st.tabs(["Job Roles & Rates", "Database & Projects", "Report Generation"])

# ========================
# JOB ROLES TAB
# ========================
with tab_roles:
    st.subheader("Manage Job Roles")
    st.markdown("Edit existing roles or add new ones. Changes apply to **new** project entries — existing projects keep their saved rates.")

    job_codes = query("SELECT * FROM job_codes ORDER BY title")

    # --- Rate Summary Table ---
    if job_codes:
        summary_df = pd.DataFrame(job_codes)
        summary_df["margin"] = summary_df["bid_rate"] - summary_df["employee_rate"]
        summary_df["margin_pct"] = (summary_df["margin"] / summary_df["bid_rate"] * 100).round(1)
        display = summary_df[["title", "bid_rate", "employee_rate", "margin", "margin_pct"]].copy()
        display.columns = ["Title", "Bid Rate", "Employee Rate", "Margin ($)", "Margin (%)"]
        st.dataframe(
            display.style.format({
                "Bid Rate": "${:,.2f}", "Employee Rate": "${:,.2f}",
                "Margin ($)": "${:,.2f}", "Margin (%)": "{:.1f}%",
            }),
            use_container_width=True, hide_index=True,
        )

        # KPI row
        avg_bid = summary_df["bid_rate"].mean()
        avg_emp = summary_df["employee_rate"].mean()
        avg_margin = summary_df["margin_pct"].mean()
        k1, k2, k3, k4 = st.columns(4)
        k1.metric("Total Roles", len(job_codes))
        k2.metric("Avg Bid Rate", f"${avg_bid:,.2f}")
        k3.metric("Avg Employee Rate", f"${avg_emp:,.2f}")
        k4.metric("Avg Margin", f"{avg_margin:.1f}%")

    st.markdown("---")

    # --- Edit / Add Section ---
    role_options = ["-- Add New Role --"] + [jc["title"] for jc in job_codes] if job_codes else ["-- Add New Role --"]
    jc_map = {jc["title"]: jc for jc in job_codes} if job_codes else {}

    # Selectbox to choose what to do
    selected_role = st.selectbox("Select a role to edit, or add a new one:", role_options, key="role_selector")

    is_adding = selected_role == "-- Add New Role --"
    editing_jc = jc_map.get(selected_role)

    with st.form("role_form"):
        st.markdown(f"### {'Add New Role' if is_adding else f'Edit: {selected_role}'}")

        fc1, fc2, fc3 = st.columns(3)
        with fc1:
            form_title = st.text_input(
                "Title",
                value="" if is_adding else editing_jc["title"],
                placeholder="e.g. Analyst 3 - SIGINT",
            )
        with fc2:
            form_bid = st.number_input(
                "Bid Rate ($/hr)",
                value=120.0 if is_adding else float(editing_jc["bid_rate"]),
                step=5.0, min_value=0.0,
            )
        with fc3:
            form_emp = st.number_input(
                "Employee Rate ($/hr)",
                value=100.0 if is_adding else float(editing_jc["employee_rate"]),
                step=5.0, min_value=0.0,
            )

        # Live margin preview
        if form_bid > 0:
            preview_margin = form_bid - form_emp
            preview_pct = (preview_margin / form_bid * 100)
            st.markdown(f"**Margin Preview:** ${preview_margin:,.2f} ({preview_pct:.1f}%)")

        btn_col1, btn_col2 = st.columns([1, 1])
        with btn_col1:
            submitted = st.form_submit_button(
                "Add Role" if is_adding else "Save Changes",
                type="primary", use_container_width=True,
            )
        with btn_col2:
            if not is_adding:
                delete_clicked = st.form_submit_button(
                    "Delete Role", use_container_width=True,
                )
            else:
                delete_clicked = False

        if submitted:
            if not form_title.strip():
                st.error("Title is required.")
            elif is_adding:
                existing = query("SELECT id FROM job_codes WHERE title = ?", (form_title.strip(),))
                if existing:
                    st.error(f"Role '{form_title}' already exists.")
                else:
                    execute(
                        "INSERT INTO job_codes (title, bid_rate, employee_rate) VALUES (?, ?, ?)",
                        (form_title.strip(), form_bid, form_emp),
                    )
                    st.toast(f"Added: {form_title}", icon="✅")
                    st.rerun()
            else:
                execute(
                    "UPDATE job_codes SET title=?, bid_rate=?, employee_rate=? WHERE id=?",
                    (form_title.strip(), form_bid, form_emp, editing_jc["id"]),
                )
                st.toast(f"Updated: {form_title}", icon="✅")
                st.rerun()

        if delete_clicked and not is_adding:
            # Check if role is used in any labor entries
            usage = query(
                "SELECT COUNT(*) as cnt FROM labor_entries WHERE job_code_id = ?",
                (editing_jc["id"],),
                fetchone=True,
            )
            if usage["cnt"] > 0:
                st.error(f"Cannot delete '{selected_role}' — it's used in {usage['cnt']} labor entry(ies). Remove those first.")
            else:
                execute("DELETE FROM job_codes WHERE id = ?", (editing_jc["id"],))
                st.toast(f"Deleted: {selected_role}", icon="🗑️")
                st.rerun()

# ========================
# DATABASE / PROJECTS TAB
# ========================
with tab_db:
    # --- DB Stats ---
    st.subheader("Database Info")

    project_count = query("SELECT COUNT(*) as cnt FROM projects", fetchone=True)["cnt"]
    labor_count = query("SELECT COUNT(*) as cnt FROM labor_entries", fetchone=True)["cnt"]
    imagery_count = query("SELECT COUNT(*) as cnt FROM imagery_orders", fetchone=True)["cnt"]
    catalog_count = query("SELECT COUNT(*) as cnt FROM imagery_catalog", fetchone=True)["cnt"]

    d1, d2, d3, d4 = st.columns(4)
    d1.metric("Projects", project_count)
    d2.metric("Labor Entries", labor_count)
    d3.metric("Imagery Orders", imagery_count)
    d4.metric("Catalog Products", catalog_count)

    db_path = Path(__file__).parent.parent / "data" / "tracker.db"
    if db_path.exists():
        size_mb = db_path.stat().st_size / (1024 * 1024)
        st.caption(f"Database: `{db_path}` | Size: {size_mb:.2f} MB")

    # --- Project Manager ---
    st.markdown("---")
    st.subheader("Project Manager")

    all_projects = query("SELECT * FROM projects ORDER BY pws_number, report_title")

    if not all_projects:
        st.info("No projects yet.")
    else:
        # PWS filter
        all_pws = sorted(set(p["pws_number"] for p in all_projects))
        selected_pws = st.selectbox("Filter by PWS", ["-- All --"] + all_pws, key="settings_pws")

        if selected_pws == "-- All --":
            filtered = all_projects
        else:
            filtered = [p for p in all_projects if p["pws_number"] == selected_pws]

        # Projects summary table
        proj_df = pd.DataFrame(filtered)
        display_cols = ["pws_number", "report_title", "status", "start_date", "end_date", "days"]
        display_df = proj_df[display_cols].copy()
        display_df.columns = ["PWS", "Title", "Status", "Start", "End", "Days"]
        st.dataframe(display_df, use_container_width=True, hide_index=True)

        # Project selector for edit/delete
        proj_map = {f"{p['pws_number']} - {p['report_title']}": p for p in filtered}
        selected_proj_name = st.selectbox(
            "Select a project to view/edit:",
            list(proj_map.keys()),
            key="settings_proj_select",
        )
        proj = proj_map[selected_proj_name]

        # Project detail card
        st.markdown(f"### {proj['report_title']}")
        info1, info2, info3 = st.columns(3)
        info1.markdown(f"**PWS:** {proj['pws_number']}")
        info2.markdown(f"**Status:** {proj['status']}")
        info3.markdown(f"**Days:** {proj['days']}")

        if proj["start_date"] or proj["end_date"]:
            st.markdown(f"**Period:** {proj['start_date'] or 'N/A'} to {proj['end_date'] or 'N/A'}")
        if proj["notes"]:
            st.markdown(f"**Notes:** {proj['notes']}")

        # Labor & imagery counts for this project
        p_labor = query(
            "SELECT COUNT(*) as cnt FROM labor_entries WHERE project_id = ?",
            (proj["id"],), fetchone=True,
        )["cnt"]
        p_imagery = query(
            "SELECT COUNT(*) as cnt FROM imagery_orders WHERE project_id = ?",
            (proj["id"],), fetchone=True,
        )["cnt"]
        st.caption(f"{p_labor} personnel entries | {p_imagery} imagery orders")

        # Quick edit form
        with st.expander("Edit Project Details"):
            with st.form("edit_proj_form"):
                ep1, ep2, ep3 = st.columns(3)
                with ep1:
                    edit_pws = st.text_input("PWS Number", value=proj["pws_number"], key="ep_pws")
                    edit_title = st.text_input("Report Title", value=proj["report_title"], key="ep_title")
                with ep2:
                    edit_status = st.selectbox(
                        "Status", ["Ongoing", "Complete"],
                        index=0 if proj["status"] == "Ongoing" else 1,
                        key="ep_status",
                    )
                    edit_days = st.number_input("Days", value=proj["days"] or 1, min_value=0, key="ep_days")
                with ep3:
                    from datetime import datetime
                    edit_start = st.date_input(
                        "Start Date",
                        value=datetime.strptime(proj["start_date"], "%Y-%m-%d").date() if proj["start_date"] else None,
                        key="ep_start",
                    )
                    edit_end = st.date_input(
                        "End Date",
                        value=datetime.strptime(proj["end_date"], "%Y-%m-%d").date() if proj["end_date"] else None,
                        key="ep_end",
                    )
                edit_notes = st.text_area("Notes", value=proj["notes"] or "", key="ep_notes")

                if st.form_submit_button("Save Changes", type="primary", use_container_width=True):
                    if not edit_pws.strip() or not edit_title.strip():
                        st.error("PWS and Title are required.")
                    else:
                        execute(
                            """UPDATE projects SET pws_number=?, report_title=?, start_date=?, end_date=?,
                               status=?, notes=?, days=? WHERE id=?""",
                            (edit_pws.strip(), edit_title.strip(), str(edit_start), str(edit_end),
                             edit_status, edit_notes or None, edit_days, proj["id"]),
                        )
                        st.toast(f"Updated: {edit_title}", icon="✅")
                        st.rerun()

        # Delete with confirmation
        with st.expander("Delete Project"):
            st.warning(f"This will permanently delete **{proj['report_title']}** and all its labor entries and imagery orders.")
            if p_labor > 0 or p_imagery > 0:
                st.markdown(f"**This will also remove {p_labor} labor entries and {p_imagery} imagery orders.**")

            if st.button("Delete This Project", key="settings_del_proj"):
                st.session_state["settings_confirm_del"] = proj["id"]

            if st.session_state.get("settings_confirm_del") == proj["id"]:
                st.error("Are you sure? This cannot be undone.")
                dc1, dc2 = st.columns(2)
                if dc1.button("Yes, Delete", key="settings_confirm_del_btn"):
                    execute("DELETE FROM labor_entries WHERE project_id = ?", (proj["id"],))
                    execute("DELETE FROM imagery_orders WHERE project_id = ?", (proj["id"],))
                    execute("DELETE FROM projects WHERE id = ?", (proj["id"],))
                    st.session_state.pop("settings_confirm_del", None)
                    st.toast(f"Deleted: {proj['report_title']}", icon="🗑️")
                    st.rerun()
                if dc2.button("Cancel", key="settings_cancel_del_btn"):
                    st.session_state.pop("settings_confirm_del", None)
                    st.rerun()

# ========================
# REPORT GENERATION TAB
# ========================
with tab_report:
    st.subheader("Generate Options Invoice")
    st.caption("Fills the standard Options Invoice template with live data from the tracker.")

    all_rpt_projects = query(
        "SELECT id, pws_number, report_title, days, is_daily_rate FROM projects ORDER BY pws_number, report_title"
    )

    if not all_rpt_projects:
        st.info("No projects yet. Create projects on the Projects page first.")
        st.stop()

    all_pws_rpt = sorted(set(p["pws_number"] for p in all_rpt_projects))

    # ── Invoice details ──────────────────────────────────────────────────────
    st.markdown("#### Invoice Details")
    d1, d2 = st.columns(2)
    with d1:
        sel_pws_rpt = st.selectbox("PWS Number", all_pws_rpt, key="rpt_pws")
        pws_projects = [p for p in all_rpt_projects if p["pws_number"] == sel_pws_rpt]
        proj_labels = [p["report_title"] for p in pws_projects]
        sel_proj_title = st.selectbox("Project", proj_labels, key="rpt_project")
        sel_proj = next(p for p in pws_projects if p["report_title"] == sel_proj_title)
        contract_event = st.text_input(
            "Contract Number \u2013 Event Name",
            value=f"PWS {sel_proj['pws_number']}",
            placeholder="e.g. 11080 - etbf / PWS 11080 / LANDAC / SPAC",
            key="rpt_contract",
        )
    with d2:
        sub_date = st.date_input("Option Submission Date", value=date.today(), key="rpt_date")
        sub_date_str = sub_date.strftime("%-d %B %Y")

    # ── Daily labour rate ────────────────────────────────────────────────────
    st.markdown("---")
    st.markdown("#### Daily Labour Rate")

    # Auto-calculate chargeable days from this project's days field
    auto_days = int(sel_proj["days"] or 0)
    chargeable_days = st.number_input(
        "Chargeable Days",
        value=auto_days,
        min_value=0,
        step=1,
        key="rpt_chargeable_days",
        help=f"Auto-filled from project days field ({auto_days}). Edit freely.",
    )

    REPORT_PW = "scott"
    unlocked = st.session_state.get("_rpt_unlocked", False)

    if not unlocked:
        pw_col, btn_col = st.columns([3, 1])
        with pw_col:
            entered_pw = st.text_input(
                "Password to view/edit daily rate", type="password", key="rpt_pw_input",
                placeholder="Enter password to unlock rate field",
            )
        with btn_col:
            st.markdown("<div style='height:28px'></div>", unsafe_allow_html=True)
            if st.button("Unlock", key="rpt_unlock_btn"):
                if entered_pw == REPORT_PW:
                    st.session_state["_rpt_unlocked"] = True
                    st.rerun()
                else:
                    st.error("Incorrect password")
        daily_rate_per_day = 3000
    else:
        rate_col, lock_col = st.columns([3, 1])
        with rate_col:
            daily_rate_per_day = st.number_input(
                "Daily Rate ($ per day)", value=3000, step=100, min_value=0, key="rpt_daily_rate",
            )
        with lock_col:
            st.markdown("<div style='height:28px'></div>", unsafe_allow_html=True)
            if st.button("Lock", key="rpt_lock_btn"):
                st.session_state.pop("_rpt_unlocked", None)
                st.rerun()

    labor_total = chargeable_days * daily_rate_per_day
    st.markdown(
        f"**Labour Total: ${labor_total:,}** "
        f"({chargeable_days} days \u00d7 ${daily_rate_per_day:,})"
    )

    # ── Imagery rows ─────────────────────────────────────────────────────────
    st.markdown("---")
    st.markdown("#### Imagery Rows")
    st.caption(
        "Auto-loaded from imagery orders with delivered shots > 0 for this project. "
        "Edit, remove, or add rows before generating."
    )

    lh1, lh2, lh3, lh4 = st.columns([3, 1, 1.5, 0.5])
    lh1.markdown("**Product / Name**")
    lh2.markdown("**Qty (delivered)**")
    lh3.markdown("**Price per unit ($)**")
    lh4.markdown("")

    db_imagery = query(
        """SELECT provider, product, shots_delivered, charge_per_shot
           FROM imagery_orders
           WHERE project_id = ? AND shots_delivered > 0
           ORDER BY provider, product""",
        (sel_proj["id"],),
    )

    # Reset imagery rows when project changes
    if st.session_state.get("_rpt_imagery_proj") != sel_proj["id"]:
        st.session_state["_rpt_imagery_proj"] = sel_proj["id"]
        st.session_state["_rpt_imagery_rows"] = [
            {
                "name": f"{r['provider']} {r['product']}",
                "qty": int(r["shots_delivered"]),
                "price": float(r["charge_per_shot"]),
            }
            for r in db_imagery
        ]

    rows = st.session_state["_rpt_imagery_rows"]
    to_delete = []

    for i, row in enumerate(rows):
        rc1, rc2, rc3, rc4 = st.columns([3, 1, 1.5, 0.5])
        with rc1:
            rows[i]["name"] = st.text_input(
                "name", value=row["name"], key=f"rpt_img_name_{i}",
                label_visibility="collapsed", placeholder="Product name",
            )
        with rc2:
            rows[i]["qty"] = st.number_input(
                "qty", value=row["qty"], min_value=0, step=1,
                key=f"rpt_img_qty_{i}", label_visibility="collapsed",
            )
        with rc3:
            rows[i]["price"] = st.number_input(
                "price", value=row["price"], min_value=0.0, step=10.0,
                key=f"rpt_img_price_{i}", label_visibility="collapsed",
            )
        with rc4:
            if st.button("\u2715", key=f"rpt_img_del_{i}", help="Remove this row"):
                to_delete.append(i)

    for i in reversed(to_delete):
        rows.pop(i)
    if to_delete:
        st.rerun()

    if st.button("+ Add Imagery Row", key="rpt_add_row"):
        rows.append({"name": "", "qty": 0, "price": 0.0})
        st.rerun()

    img_total = sum(r["qty"] * r["price"] for r in rows)
    st.markdown(f"**Imagery Total: ${img_total:,.2f}**")

    # ── Output format & generate ─────────────────────────────────────────────
    st.markdown("---")
    st.markdown("#### Generate Report")
    out_format = st.radio(
        "Output Format",
        ["Word (.docx)", "PDF (.pdf)", "Both"],
        horizontal=True,
        key="rpt_format",
    )

    if st.button("Generate Report", type="primary", key="rpt_generate"):
        labor_row_data = {
            "name": "Stone Harp Daily Labor Rate",
            "qty": chargeable_days,
            "price": daily_rate_per_day,
            "desc": f"Quantity * price = ${chargeable_days * daily_rate_per_day:,}",
        }
        img_rows_data = [
            {
                "name": r["name"],
                "qty": r["qty"],
                "price": r["price"],
                "desc": f"Quantity * price = ${r['qty'] * r['price']:,.2f}",
            }
            for r in rows
            if r["name"].strip()
        ]
        hdr = {
            "contract_event": contract_event,
            "submission_date": sub_date_str,
        }
        safe_title = sel_proj["report_title"].replace(" ", "_")[:30]
        filename_base = f"Options_Invoice_{safe_title}_{sub_date.strftime('%Y%m%d')}"

        dl_col1, dl_col2 = st.columns(2)
        if out_format in ("Word (.docx)", "Both"):
            try:
                docx_bytes = _generate_docx(hdr, labor_row_data, img_rows_data)
                dl_col1.download_button(
                    "Download Word (.docx)",
                    data=docx_bytes,
                    file_name=f"{filename_base}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            except Exception as e:
                st.error(f"DOCX generation failed: {e}")

        if out_format in ("PDF (.pdf)", "Both"):
            try:
                pdf_bytes = _generate_pdf(hdr, labor_row_data, img_rows_data)
                dl_col2.download_button(
                    "Download PDF (.pdf)",
                    data=pdf_bytes,
                    file_name=f"{filename_base}.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                )
            except Exception as e:
                st.error(f"PDF generation failed: {e}")
